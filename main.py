import os
import requests
import re
import time
import traceback
import tempfile
import argparse
import hashlib
from urllib.parse import urlparse
from bs4 import BeautifulSoup, NavigableString
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from bypass_wechat_limitations import fetch_wechat_article

def is_valid_url(url):
    """Check if the URL is valid"""
    try:
        result = urlparse(url)
        return all([result.scheme, result.netloc])
    except:
        return False

def download_image(img_url, temp_dir):
    """Download image from URL and save to temp directory"""
    try:
        # Ensure the URL is absolute
        if not img_url.startswith(('http://', 'https://')):
            if img_url.startswith('//'):
                img_url = 'https:' + img_url
            else:
                print(f"Skipping image with invalid URL: {img_url}")
                return None

        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36',
            'Referer': 'https://mp.weixin.qq.com/'
        }
        response = requests.get(img_url, headers=headers, stream=True, timeout=10)
        if response.status_code == 200:
            # Create a unique filename based on URL content
            img_hash = hashlib.md5(img_url.encode()).hexdigest()[:10]
            img_filename = f"img_{img_hash}_{int(time.time())}.jpg"
            img_path = os.path.join(temp_dir, img_filename)
            with open(img_path, 'wb') as f:
                for chunk in response.iter_content(1024):
                    f.write(chunk)
            # Wait a moment to avoid being blocked
            time.sleep(0.2)
            return img_path
    except Exception as e:
        print(f"Failed to download image: {e} - URL: {img_url}")
    return None

def is_significant_text(text):
    """Check if text is significant (not just whitespace or short)"""
    if not text:
        return False
    
    # Clean the text
    text = text.strip()
    
    # Check if it contains actual content
    if len(text) < 2:  # Too short to be meaningful
        return False
        
    # Check if it's just punctuation or special characters
    if re.match(r'^[.,;:!?\-_=+*&^%$#@<>()[\]{}|~\'"]+$', text):
        return False
        
    return True

def extract_content_tree(element, processed_elements=None, depth=0, parent_path="", max_depth=20):
    """
    Extract content as a hierarchical tree to maintain structure and prevent duplicates
    Returns a list of dictionaries with content and metadata
    """
    if processed_elements is None:
        processed_elements = set()
        
    if depth > max_depth:  # Prevent infinite recursion
        return []
        
    if not element or id(element) in processed_elements:
        return []
    
    # Mark this element as processed
    processed_elements.add(id(element))
    
    current_path = f"{parent_path}/{element.name}[{id(element)}]"
    result = []
    
    # Special case for comment objects
    if isinstance(element, NavigableString):
        if is_significant_text(str(element)):
            return [{'type': 'text', 'content': str(element).strip(), 'path': current_path}]
        return []
    
    # Check if this is an image element
    if element.name == 'img':
        img_url = None
        for attr in ['data-src', 'src', 'data-url', 'data-backh-src']:
            if element.get(attr):
                img_url = element.get(attr)
                break
        
        if img_url:
            return [{'type': 'image', 'url': img_url, 'alt': element.get('alt', ''), 'path': current_path}]
        return []
    
    # Handle headers specially
    if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
        text = element.get_text().strip()
        if text:
            level = int(element.name[1])
            return [{'type': 'heading', 'level': level, 'content': text, 'path': current_path}]
    
    # For <br> tags, add a line break
    if element.name == 'br':
        return [{'type': 'break', 'path': current_path}]
    
    # For list items, handle specially
    if element.name == 'li':
        text = element.get_text().strip()
        list_type = 'bullet'  # Default
        
        # Find parent list type
        parent = element.parent
        if parent and parent.name == 'ol':
            list_type = 'numbered'
            
        if text:
            return [{'type': 'list_item', 'list_type': list_type, 'content': text, 'path': current_path}]
    
    # Handle paragraph-like elements
    is_paragraph = element.name in ['p', 'div', 'section', 'article']
    
    # Process direct text content of this element (ignoring child elements)
    direct_text = ''.join(str(child) for child in element.contents 
                          if isinstance(child, NavigableString)).strip()
                          
    # Add direct text if it exists and is not already in a paragraph element
    if is_significant_text(direct_text):
        if is_paragraph:
            result.append({'type': 'paragraph', 'content': direct_text, 'path': current_path})
        else:
            result.append({'type': 'text', 'content': direct_text, 'path': current_path})
    
    # Recursively process children
    for child in element.children:
        if not isinstance(child, NavigableString) or is_significant_text(str(child)):
            child_content = extract_content_tree(child, processed_elements, depth + 1, current_path, max_depth)
            result.extend(child_content)
    
    return result

def post_process_content(content_tree):
    """
    Post-process the content tree to:
    - Combine adjacent text fragments
    - Handle special cases
    - Fix formatting
    """
    if not content_tree:
        return []
    
    result = []
    current_paragraph = None
    list_items_buffer = []
    current_list_type = None
    
    for item in content_tree:
        item_type = item.get('type')
        
        # Handle list items: group them together
        if item_type == 'list_item':
            if current_paragraph:
                result.append(current_paragraph)
                current_paragraph = None
                
            # Check if we're continuing the same list or starting a new one
            if not list_items_buffer or current_list_type == item.get('list_type'):
                list_items_buffer.append(item.get('content'))
                current_list_type = item.get('list_type')
            else:
                # Finish the current list and start a new one
                result.append({
                    'type': 'list',
                    'list_type': current_list_type,
                    'items': list_items_buffer
                })
                list_items_buffer = [item.get('content')]
                current_list_type = item.get('list_type')
            continue
        
        # When we encounter a non-list item, finish any current list
        if list_items_buffer:
            result.append({
                'type': 'list',
                'list_type': current_list_type,
                'items': list_items_buffer
            })
            list_items_buffer = []
            current_list_type = None
        
        # Handle regular text/paragraphs
        if item_type in ['text', 'paragraph']:
            if current_paragraph:
                # If the previous item is also text, append to it with a space
                if item.get('content').strip():
                    current_paragraph['content'] += ' ' + item.get('content').strip()
            else:
                current_paragraph = {
                    'type': 'paragraph',
                    'content': item.get('content').strip()
                }
        # Handle line breaks
        elif item_type == 'break':
            if current_paragraph:
                current_paragraph['content'] += '\n'
        # Handle headings and images - they break paragraphs
        elif item_type in ['heading', 'image']:
            if current_paragraph:
                result.append(current_paragraph)
                current_paragraph = None
            result.append(item)
    
    # Add any final paragraph or list
    if current_paragraph:
        result.append(current_paragraph)
        
    if list_items_buffer:
        result.append({
            'type': 'list',
            'list_type': current_list_type,
            'items': list_items_buffer
        })
    
    # Clean up paragraph content
    for item in result:
        if item.get('type') == 'paragraph':
            # Replace multiple spaces with single space
            item['content'] = re.sub(r'\s+', ' ', item['content'])
            # Handle paragraph breaks
            item['content'] = item['content'].replace('\n\n', '\n')
            # Final trim
            item['content'] = item['content'].strip()
    
    return result

def process_wechat_article(html_content):
    """
    Process the WeChat article HTML to extract structured content
    Returns a dictionary with title, author, and content blocks
    """
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Find all possible content containers
    content_selectors = [
        'div.rich_media_content', 
        'div#js_content',
        'div.article-content',
        'div.content-article',
        'div.wx-article-content'
    ]
    
    content_div = None
    for selector in content_selectors:
        elements = soup.select(selector)
        if elements:
            content_div = elements[0]
            print(f"Found content using selector: {selector}")
            break
    
    # If content container not found, try a broader approach
    if not content_div:
        print("Could not find main content container, trying backup method...")
        content_candidates = soup.find_all('div', class_=re.compile(r'(article|content|text)'))
        if content_candidates:
            content_div = max(content_candidates, key=lambda x: len(x.get_text()))
            print("Using backup content container")
    
    # Extract title using multiple potential selectors
    title = None
    title_selectors = [
        'h1.rich_media_title',
        'h1#activity-name',
        'h1.activity-name',
        'h2.rich_media_title',
        'h1.title',
        'meta[property="og:title"]'
    ]
    
    for selector in title_selectors:
        elements = soup.select(selector)
        if elements:
            if selector.startswith('meta'):
                title = elements[0].get('content', '')
            else:
                title = elements[0].get_text().strip()
            if title:
                print(f"Found title: {title}")
                break
    
    if not title:
        print("Could not find title using selectors, searching in document...")
        h1_elements = soup.find_all('h1')
        if h1_elements:
            title = h1_elements[0].get_text().strip()
            print(f"Found title from h1: {title}")
        else:
            # Look for large text at the beginning that might be the title
            large_text_elements = soup.find_all(['h2', 'div', 'p'], class_=re.compile(r'(title|headline)'))
            if large_text_elements:
                title = large_text_elements[0].get_text().strip()
                print(f"Found title from large text: {title}")
            else:
                title = "WeChat Article"
                print("Using default title")
    
    # Extract content using the improved hierarchical method
    content_tree = []
    if content_div:
        print("Extracting content tree from main content div...")
        content_tree = extract_content_tree(content_div)
    else:
        print("No content div found, extracting from body...")
        content_tree = extract_content_tree(soup.body)
    
    # Post-process the content tree to combine text and fix formatting
    content_blocks = post_process_content(content_tree)
    
    return {
        'title': title,
        'blocks': content_blocks
    }

def add_paragraph_with_formatting(doc, text):
    """Add paragraph to document with proper formatting for CJK text"""
    paragraph = doc.add_paragraph()
    
    # Split by newlines to handle manual line breaks
    parts = text.split('\n')
    
    for i, part in enumerate(parts):
        if i > 0:  # Add line break between parts
            paragraph.add_run().add_break()
        
        # Add text with proper formatting
        run = paragraph.add_run(part)
        run.font.name = 'Microsoft YaHei'
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    return paragraph

def wechat_to_docx(url, output_path=None):
    """Convert WeChat article to docx document"""
    if not is_valid_url(url):
        print("Invalid URL provided")
        return False
    
    try:
        print(f"Fetching article from {url}...")
        
        # Use enhanced fetching method
        html_content = fetch_wechat_article(url)
        
        if not html_content:
            print("Failed to retrieve the article content")
            return False
        
        # Save HTML for debugging
        with open('debug_html.html', 'w', encoding='utf-8') as f:
            f.write(html_content)
                
        print("Saved HTML to debug_html.html for debugging")
        
        # Process the article
        print("Processing article...")
        article_data = process_wechat_article(html_content)
        
        # Create document
        doc = Document()
        
        # Set document properties for better CJK character support
        doc.styles['Normal'].font.name = 'Microsoft YaHei'
        doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        
        # Add title
        title = article_data['title']
        heading = doc.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Create temporary directory for images
        temp_dir = tempfile.mkdtemp()
        print(f"Created temporary directory for images: {temp_dir}")
        
        # Process content blocks
        print(f"Adding {len(article_data['blocks'])} content blocks to document...")
        for i, block in enumerate(article_data['blocks']):
            block_type = block.get('type', '')
            
            if block_type == 'image':
                img_url = block['url']
                print(f"Processing image: {img_url[:50]}...")
                img_path = download_image(img_url, temp_dir)
                if img_path:
                    try:
                        # Add a paragraph break before image for better spacing
                        doc.add_paragraph()
                        
                        # Add image centered
                        p = doc.add_paragraph()
                        r = p.add_run()
                        r.add_picture(img_path, width=Inches(5.5))
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        
                        # Add another paragraph break after image
                        doc.add_paragraph()
                        print("Added image to document")
                    except Exception as img_ex:
                        print(f"Error adding image to document: {img_ex}")
                        
            elif block_type == 'heading':
                level = block.get('level', 2)
                text = block.get('content', '').strip()
                if text:
                    # Make sure level is between 1-5
                    adjusted_level = max(1, min(5, level + 1))  # Offset by 1 since title is level 1
                    h = doc.add_heading(text, level=adjusted_level)
                    print(f"Added heading (L{adjusted_level}): {text[:30]}...")
                    
            elif block_type == 'paragraph':
                text = block.get('content', '').strip()
                if text:
                    # Add paragraph with proper formatting
                    add_paragraph_with_formatting(doc, text)
                    print(f"Added paragraph: {text[:30]}...")
                    
            elif block_type == 'list':
                items = block.get('items', [])
                list_type = block.get('list_type', 'bullet')
                
                for item in items:
                    if item.strip():
                        p = doc.add_paragraph(style='List Bullet' if list_type == 'bullet' else 'List Number')
                        p.add_run(item.strip())
                        print(f"Added list item: {item[:30]}...")
                        
        # Save document
        if not output_path:
            # Sanitize filename
            safe_title = re.sub(r'[\\/*?:"<>|]', "_", title)
            output_path = f"{safe_title[:50].replace(' ', '_')}.docx"
            
        print(f"Saving document to {output_path}...")
        doc.save(output_path)
        print(f"Document saved as: {output_path}")
        
        return True
    except Exception as e:
            print(f"Error converting article: {e}")
            print("Detailed traceback:")
            traceback.print_exc()
            return False
        
def main():
    parser = argparse.ArgumentParser(description='Convert WeChat article to Word document')
    parser.add_argument('url', help='URL of the WeChat article')
    parser.add_argument('--output', '-o', help='Output document path')
    args = parser.parse_args()
    success = wechat_to_docx(args.url, args.output)
    if success:
        print("Conversion completed successfully!")
    else:
        print("Conversion failed.")
    
if __name__ == "__main__":
        main()
